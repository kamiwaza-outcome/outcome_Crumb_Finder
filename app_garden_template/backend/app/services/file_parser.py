"""
File parsing framework for App Garden applications.

This module provides a flexible system for parsing various file formats
with built-in security checks and validation.
"""

import io
import csv
import json
import logging
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any, Dict, List, Optional, BinaryIO

# Document parsing libraries (optional dependencies)
try:
    import pypdf
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import markdown
    HAS_MARKDOWN = True
except ImportError:
    HAS_MARKDOWN = False

try:
    import openpyxl
    HAS_EXCEL = True
except ImportError:
    HAS_EXCEL = False

from ..core.errors import FileError, FileSizeError, FileTypeError
from ..models.base import FileInfo

logger = logging.getLogger(__name__)


class BaseFileParser(ABC):
    """
    Abstract base class for file parsers.
    
    Provides common functionality for validation, security checks,
    and error handling.
    """
    
    # File format configuration
    SUPPORTED_EXTENSIONS: List[str] = []
    SUPPORTED_MIME_TYPES: List[str] = []
    MAX_FILE_SIZE: int = 10 * 1024 * 1024  # 10MB default
    
    def __init__(self, max_file_size: Optional[int] = None):
        """
        Initialize parser with optional custom file size limit.
        
        Args:
            max_file_size: Maximum file size in bytes (overrides class default)
        """
        if max_file_size:
            self.MAX_FILE_SIZE = max_file_size
        self.logger = logging.getLogger(self.__class__.__name__)
    
    def validate_file(self, filename: str, file_size: int, content_type: Optional[str] = None) -> FileInfo:
        """
        Validate file before parsing.
        
        Args:
            filename: Name of the file
            file_size: Size in bytes
            content_type: MIME type if available
            
        Returns:
            FileInfo object with validated metadata
            
        Raises:
            FileSizeError: If file exceeds size limit
            FileTypeError: If file type is not supported
        """
        # Check file size
        if file_size > self.MAX_FILE_SIZE:
            raise FileSizeError(
                filename=filename,
                size=file_size,
                max_size=self.MAX_FILE_SIZE
            )
        
        # Check file extension
        file_ext = Path(filename).suffix.lower()
        if self.SUPPORTED_EXTENSIONS and file_ext not in self.SUPPORTED_EXTENSIONS:
            raise FileTypeError(
                filename=filename,
                file_type=file_ext,
                allowed_types=self.SUPPORTED_EXTENSIONS
            )
        
        # Check MIME type if provided
        if content_type and self.SUPPORTED_MIME_TYPES:
            if content_type not in self.SUPPORTED_MIME_TYPES:
                raise FileTypeError(
                    filename=filename,
                    file_type=content_type,
                    allowed_types=self.SUPPORTED_MIME_TYPES
                )
        
        return FileInfo(
            filename=filename,
            content_type=content_type or "application/octet-stream",
            size=file_size
        )
    
    @abstractmethod
    async def parse(self, file_content: bytes, file_info: FileInfo) -> Any:
        """
        Parse file content.
        
        Args:
            file_content: Raw file bytes
            file_info: Validated file information
            
        Returns:
            Parsed content (format depends on parser implementation)
        """
        pass
    
    async def parse_file(self, filename: str, file_content: bytes, content_type: Optional[str] = None) -> Any:
        """
        Validate and parse a file.
        
        Args:
            filename: Name of the file
            file_content: Raw file bytes
            content_type: Optional MIME type
            
        Returns:
            Parsed content
        """
        file_info = self.validate_file(filename, len(file_content), content_type)
        
        self.logger.info(
            f"Parsing file: {filename}",
            extra={
                "filename": filename,
                "size": len(file_content),
                "content_type": content_type
            }
        )
        
        try:
            return await self.parse(file_content, file_info)
        except Exception as e:
            self.logger.error(f"Failed to parse file {filename}: {e}")
            raise FileError(
                message=f"Failed to parse file: {str(e)}",
                filename=filename
            )


class TextFileParser(BaseFileParser):
    """Parser for plain text files."""
    
    SUPPORTED_EXTENSIONS = ['.txt', '.text', '.log']
    SUPPORTED_MIME_TYPES = ['text/plain']
    
    async def parse(self, file_content: bytes, file_info: FileInfo) -> str:
        """Parse text file content."""
        try:
            # Try UTF-8 first, then fallback to other encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use UTF-8 with error handling
            return file_content.decode('utf-8', errors='replace')
        
        except Exception as e:
            raise FileError(f"Failed to decode text file: {e}", file_info.filename)


class JSONFileParser(BaseFileParser):
    """Parser for JSON files."""
    
    SUPPORTED_EXTENSIONS = ['.json']
    SUPPORTED_MIME_TYPES = ['application/json']
    
    async def parse(self, file_content: bytes, file_info: FileInfo) -> Dict[str, Any]:
        """Parse JSON file content."""
        try:
            text = file_content.decode('utf-8')
            return json.loads(text)
        except json.JSONDecodeError as e:
            raise FileError(
                f"Invalid JSON format: {e}",
                file_info.filename,
                details={"line": e.lineno, "column": e.colno}
            )


class CSVFileParser(BaseFileParser):
    """Parser for CSV files."""
    
    SUPPORTED_EXTENSIONS = ['.csv']
    SUPPORTED_MIME_TYPES = ['text/csv', 'application/csv']
    
    def __init__(self, max_file_size: Optional[int] = None, delimiter: str = ','):
        super().__init__(max_file_size)
        self.delimiter = delimiter
    
    async def parse(self, file_content: bytes, file_info: FileInfo) -> List[Dict[str, Any]]:
        """Parse CSV file content."""
        try:
            text = file_content.decode('utf-8')
            reader = csv.DictReader(io.StringIO(text), delimiter=self.delimiter)
            return list(reader)
        except Exception as e:
            raise FileError(f"Failed to parse CSV: {e}", file_info.filename)


class PDFFileParser(BaseFileParser):
    """Parser for PDF files (requires pypdf)."""
    
    SUPPORTED_EXTENSIONS = ['.pdf']
    SUPPORTED_MIME_TYPES = ['application/pdf']
    
    def __init__(self, max_file_size: Optional[int] = None):
        super().__init__(max_file_size)
        if not HAS_PDF:
            raise ImportError(
                "PDF parsing requires pypdf. Install with: pip install pypdf"
            )
    
    async def parse(self, file_content: bytes, file_info: FileInfo) -> str:
        """Extract text from PDF file."""
        try:
            pdf_reader = pypdf.PdfReader(io.BytesIO(file_content))
            
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text.strip():
                    text_parts.append(f"--- Page {page_num + 1} ---\n{text}")
            
            return "\n\n".join(text_parts)
        
        except Exception as e:
            raise FileError(f"Failed to parse PDF: {e}", file_info.filename)


class DOCXFileParser(BaseFileParser):
    """Parser for DOCX files (requires python-docx)."""
    
    SUPPORTED_EXTENSIONS = ['.docx']
    SUPPORTED_MIME_TYPES = [
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    ]
    
    def __init__(self, max_file_size: Optional[int] = None):
        super().__init__(max_file_size)
        if not HAS_DOCX:
            raise ImportError(
                "DOCX parsing requires python-docx. Install with: pip install python-docx"
            )
    
    async def parse(self, file_content: bytes, file_info: FileInfo) -> str:
        """Extract text from DOCX file."""
        try:
            doc = docx.Document(io.BytesIO(file_content))
            
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            return "\n\n".join(paragraphs)
        
        except Exception as e:
            raise FileError(f"Failed to parse DOCX: {e}", file_info.filename)


class MarkdownFileParser(BaseFileParser):
    """Parser for Markdown files (with optional HTML conversion)."""
    
    SUPPORTED_EXTENSIONS = ['.md', '.markdown']
    SUPPORTED_MIME_TYPES = ['text/markdown', 'text/x-markdown']
    
    def __init__(self, max_file_size: Optional[int] = None, convert_to_html: bool = False):
        super().__init__(max_file_size)
        self.convert_to_html = convert_to_html
    
    async def parse(self, file_content: bytes, file_info: FileInfo) -> str:
        """Parse Markdown file content."""
        try:
            text = file_content.decode('utf-8')
            
            if self.convert_to_html and HAS_MARKDOWN:
                return markdown.markdown(text)
            else:
                return text
        
        except Exception as e:
            raise FileError(f"Failed to parse Markdown: {e}", file_info.filename)


class ExcelFileParser(BaseFileParser):
    """Parser for Excel files (requires openpyxl)."""
    
    SUPPORTED_EXTENSIONS = ['.xlsx', '.xlsm']
    SUPPORTED_MIME_TYPES = [
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'application/vnd.ms-excel.sheet.macroEnabled.12'
    ]
    
    def __init__(self, max_file_size: Optional[int] = None):
        super().__init__(max_file_size)
        if not HAS_EXCEL:
            raise ImportError(
                "Excel parsing requires openpyxl. Install with: pip install openpyxl"
            )
    
    async def parse(self, file_content: bytes, file_info: FileInfo) -> Dict[str, List[List[Any]]]:
        """Parse Excel file content."""
        try:
            workbook = openpyxl.load_workbook(io.BytesIO(file_content), read_only=True)
            
            result = {}
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # Extract all data from sheet
                sheet_data = []
                for row in sheet.iter_rows(values_only=True):
                    # Filter out empty rows
                    if any(cell is not None for cell in row):
                        sheet_data.append(list(row))
                
                if sheet_data:
                    result[sheet_name] = sheet_data
            
            workbook.close()
            return result
        
        except Exception as e:
            raise FileError(f"Failed to parse Excel: {e}", file_info.filename)


class FileParserFactory:
    """
    Factory for creating appropriate file parsers based on file type.
    
    This allows dynamic parser selection and easy extension with new formats.
    """
    
    def __init__(self):
        self._parsers: Dict[str, type[BaseFileParser]] = {
            '.txt': TextFileParser,
            '.text': TextFileParser,
            '.log': TextFileParser,
            '.json': JSONFileParser,
            '.csv': CSVFileParser,
        }
        
        # Register optional parsers if dependencies are available
        if HAS_PDF:
            self._parsers['.pdf'] = PDFFileParser
        
        if HAS_DOCX:
            self._parsers['.docx'] = DOCXFileParser
        
        if HAS_MARKDOWN:
            self._parsers['.md'] = MarkdownFileParser
            self._parsers['.markdown'] = MarkdownFileParser
        
        if HAS_EXCEL:
            self._parsers['.xlsx'] = ExcelFileParser
            self._parsers['.xlsm'] = ExcelFileParser
    
    def register_parser(self, extension: str, parser_class: type[BaseFileParser]):
        """Register a custom parser for a file extension."""
        self._parsers[extension.lower()] = parser_class
    
    def get_parser(self, filename: str, **kwargs) -> BaseFileParser:
        """
        Get appropriate parser for a file.
        
        Args:
            filename: Name of the file
            **kwargs: Additional arguments for parser initialization
            
        Returns:
            Initialized parser instance
            
        Raises:
            FileTypeError: If no parser is available for the file type
        """
        ext = Path(filename).suffix.lower()
        
        if ext not in self._parsers:
            raise FileTypeError(
                filename=filename,
                file_type=ext,
                allowed_types=list(self._parsers.keys())
            )
        
        parser_class = self._parsers[ext]
        return parser_class(**kwargs)
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions."""
        return list(self._parsers.keys())


# Global factory instance
file_parser_factory = FileParserFactory()